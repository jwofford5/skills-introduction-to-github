import subprocess
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x40, 0x40, 0x40)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)
    section.different_first_page_header_footer = True  # no header on page 1

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)


def set_font(run, size=11, bold=False, italic=False, color=BLACK, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_colored_rule(doc, color_hex='1F3864', thickness='12'):
    """Full-width colored horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), thickness)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_name_block(doc, name, email, phone, location):
    add_colored_rule(doc, thickness='18')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(name)
    set_font(r, size=20, bold=True, color=NAVY)

    # Contact line with clickable email
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    add_hyperlink(p2, email, f'mailto:{email}')
    sep = p2.add_run('  •  ' + phone + '  •  ' + location)
    set_font(sep, size=10, color=DARK_GRAY)

    add_colored_rule(doc, thickness='18')


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper())
    set_font(r, size=10.5, bold=True, color=NAVY)
    add_colored_rule(doc, thickness='6')
    return p


def add_role(doc, title, org_date):
    """Bold title on line 1, italic org + date on line 2."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(7)
    p1.paragraph_format.space_after = Pt(1)
    r = p1.add_run(title)
    set_font(r, size=11, bold=True, color=BLACK)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(org_date)
    set_font(r2, size=10.5, italic=True, color=DARK_GRAY)
    return p2


def add_credential(doc, credential, org_date, note=None):
    """For education entries."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(7)
    p1.paragraph_format.space_after = Pt(1)
    r = p1.add_run(credential)
    set_font(r, size=11, bold=True, color=BLACK)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(1 if note else 2)
    r2 = p2.add_run(org_date)
    set_font(r2, size=10.5, italic=True, color=DARK_GRAY)

    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(2)
        r3 = p3.add_run(note)
        set_font(r3, size=10.5, color=DARK_GRAY)
    return p2


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)

    bullet_run = p.add_run('•  ')
    set_font(bullet_run, size=10, color=NAVY)

    text_run = p.add_run(text)
    set_font(text_run, size=10.5, color=BLACK)
    return p


def add_para(doc, text, size=10.5, space_before=2, space_after=4, color=BLACK, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, color=color, italic=italic)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=10.5, italic=True, color=DARK_GRAY)
    return p


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    # Style the run directly since 'Hyperlink' style may not exist
    run = paragraph.add_run('')
    for child in hyperlink:
        if child.tag.endswith('}r'):
            r_elem = child
            rPr2 = r_elem.get_or_add_rPr() if hasattr(r_elem, 'get_or_add_rPr') else OxmlElement('w:rPr')
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), '1F3864')
            rPr2.append(color_elem)
            u_elem = OxmlElement('w:u')
            u_elem.set(qn('w:val'), 'single')
            rPr2.append(u_elem)
    return hyperlink


def add_running_header(doc, name):
    """Add name + page number to the non-first-page header."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)

        name_run = p.add_run(name + '  |  ')
        set_font(name_run, size=9, color=DARK_GRAY, italic=True)

        # Page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        page_run.append(rPr)
        page_run.append(fldChar1)
        page_run.append(instrText)
        page_run.append(fldChar2)
        p._p.append(page_run)

        # Style the page number
        num_run = p.add_run()
        set_font(num_run, size=9, color=DARK_GRAY)


def export_pdf(docx_path):
    """Convert .docx to .pdf using LibreOffice headless. Skips gracefully if unavailable."""
    output_dir = os.path.dirname(os.path.abspath(docx_path))
    try:
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, docx_path
        ], capture_output=True, timeout=30)
        pdf_name = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        pdf_path = os.path.join(output_dir, pdf_name)
        if os.path.exists(pdf_path):
            return pdf_path
        print('Note: PDF export unavailable in this environment.')
        print('To create a PDF: open the .docx in Word → File → Save As → PDF,')
        print('or upload to Google Docs → File → Download → PDF Document.')
        return None
    except Exception as e:
        print(f'Note: PDF export skipped ({e}).')
        print('To create a PDF: open the .docx in Word → File → Save As → PDF.')
        return None


# ── NAME & CONTACT ───────────────────────────────────────────────
add_name_block(
    doc,
    'Chester Lee McMillion',
    email='lj.mcmillion@icloud.com',
    phone='714.642.1313',
    location='Los Angeles, California'
)

# ── RESEARCH INTERESTS ───────────────────────────────────────────
add_section_header(doc, 'Research Interests')
add_para(doc,
    'Counterterrorism policy and practice; transatlantic security cooperation; '
    'critical incident command and interagency coordination; comparative police '
    'and military special operations doctrine; security governance for major '
    'international events.',
    space_before=5)

# ── EDUCATION ────────────────────────────────────────────────────
add_section_header(doc, 'Education')

add_credential(doc,
    'Fulbright Scholar / Research Fellowship',
    'U.S. Department of State  •  London, England  •  September 2011 – March 2012',
    'Conducted independent research on counterterrorism and critical incident response '
    'throughout the United Kingdom. Embedded with counterterror police and military special '
    'forces units. Sponsored by the U.S. Department of State Bureau of Educational and Cultural Affairs.')

add_credential(doc,
    'Foundation for Senior Leaders',
    'National Police College, Bramshill, England  •  January 2012',
    'Executive Module, Business Module, and Professional Policing Module')

add_credential(doc,
    'Master of Public Administration (MPA)',
    'University of Southern California, Los Angeles, CA  •  December 2002',
    'Certificate of Merit, Outstanding Master’s Graduate Student  •  Pi Alpha Alpha National Honor Society')

add_credential(doc,
    'Bachelor of Arts, Political Science',
    'Union Institute, Los Angeles, CA  •  September 1993',
    'Emphasis in International Security and International Relations')

add_credential(doc,
    'Delinquency Control Institute, 110th Class',
    'University of Southern California  •  February 2000')

add_credential(doc,
    'Instructor Development Course',
    'University of California, Los Angeles  •  August 1993')

# ── INTERNATIONAL AFFAIRS & RESEARCH ────────────────────────────
add_section_header(doc, 'International Affairs and Research Experience')

add_role(doc,
    'Foreign Liaison and International Security Advisor',
    'Los Angeles Police Department  •  1996 – Present')
add_para(doc,
    'Sustained, senior-level engagement with foreign law enforcement, military, and government '
    'counterparts across the United Kingdom, France, Norway, and the United Arab Emirates, '
    'encompassing counterterrorism best practices, threat trend analysis, tactical doctrine '
    'exchange, policy coordination, and diplomatic response protocols.',
    space_before=3, space_after=2)
add_label(doc, 'Selected contributions:')
add_bullet(doc, 'Organized and led security planning and major incident debriefs for the London Olympics (2012), Paris Olympics (2024), and the upcoming Los Angeles Olympics (2028), coordinating across national and international stakeholders.')
add_bullet(doc, 'Facilitated bilateral personnel exchanges between LAPD SWAT and allied special operations units to advance shared doctrine and interoperability.')
add_bullet(doc, 'Provided leadership development training in critical incident command and interagency coordination to international partners.')
add_bullet(doc, 'Recognized as an International Exchange Alumni by the U.S. Department of State for contributions to public safety diplomacy.')

# ── PROFESSIONAL RESEARCH & ANALYSIS ────────────────────────────
add_section_header(doc, 'Professional Research and Analysis')

add_role(doc,
    'Expert Witness and Law Enforcement Policy Analyst',
    'Independent Consultant  •  2008 – Present')
add_para(doc,
    'Retained by law enforcement agencies and legal counsel to provide expert analysis and '
    'testimony in civil litigation. Each engagement requires in-depth examination of agency '
    'policies, incident reconstruction, and assessment of compliance with applicable case law '
    'and professional standards, synthesized into formal Federal Rule 26 expert reports.',
    space_before=3, space_after=2)
add_bullet(doc, 'Research encompasses use-of-force doctrine, tactical decision-making, pursuit policy, officer conduct, and supervisory responsibility.')
add_bullet(doc, 'Findings translated into structured written analysis accessible to legal, judicial, and lay audiences.')

# ── TEACHING EXPERIENCE ──────────────────────────────────────────
add_section_header(doc, 'Teaching Experience')

add_role(doc, 'Instructor', 'National Tactical Officers Association (NTOA)')
add_bullet(doc, 'Basic and Advanced SWAT Tactics; Crisis Negotiations; SWAT Team Leadership')

add_role(doc, 'Instructor', 'California Association of Tactical Officers (CATO)')
add_bullet(doc, 'SWAT Team Leadership')

add_role(doc, 'Instructor', 'Blackwater Training, USA')
add_bullet(doc, 'Basic and Advanced SWAT Tactics; Crisis Negotiations')

add_role(doc, 'Recruit and In-Service Instructor',
    'LAPD Training Division, Firearms Training Unit  •  1993 – 1994')

# ── PROFESSIONAL EXPERIENCE ──────────────────────────────────────
add_section_header(doc, 'Professional Experience')

add_para(doc,
    'Los Angeles Police Department  •  1988 – Present  •  '
    'POST Certificates: Advanced, Supervisory, and Management',
    size=10.5, italic=True, color=DARK_GRAY, space_before=5, space_after=2)

add_role(doc, 'Lieutenant II, SWAT Team Commander',
    'Metropolitan Division, Special Weapons and Tactics  •  2015 – Present')
add_bullet(doc, 'Command one of two leadership positions for LAPD SWAT — a 68-person, full-time unit averaging 150 deployments annually, including armed barricade, hostage rescue, and high-risk warrant operations.')
add_bullet(doc, 'Evaluate and approve all response tactics; maintain accountability to chain of command for deployment decisions consistent with contemporary use-of-force standards.')
add_bullet(doc, 'Over a 10-year period, 92% of approximately 1,500 incidents were resolved without use of force; only 1.48% required deadly force — reflecting institutional commitment to de-escalation.')
add_bullet(doc, 'Manage personnel development, misconduct investigations, inner-platoon promotions, and multi-source budget administration including city allocations, grants, and donations.')

add_role(doc, 'Lieutenant II, Gang Impact Team Commander',
    '77th Street Area  •  2013 – 2015')
add_bullet(doc, 'Led integrated unit comprising Gang Enforcement, Gang Investigations (detectives), and Narcotics Enforcement details; coordinated with local, state, and federal agencies.')
add_bullet(doc, 'Served as Acting Patrol Commanding Officer in the absence of the Division Captain.')

add_role(doc, 'Lieutenant I, Patrol Watch Commander',
    'Newton Patrol Division  •  2012 – 2013')
add_bullet(doc, 'Incident commander for officer-involved shootings, pursuits, and large-scale tactical situations; administrative oversight for uses of force and risk management.')

add_role(doc, 'Sergeant II, SWAT Squad Leader',
    'Metropolitan Division, Special Weapons and Tactics  •  2009 – 2012')
add_bullet(doc, 'Supervised specialty cadres: snipers, lead climbers, tactical waterborne divers, firearms instructors, and crisis negotiators.')

add_role(doc, 'Sergeant I / SWAT Tactical Waterborne Diving Supervisor',
    '77th Street Patrol Division  •  2007 – 2008')
add_bullet(doc, 'Sole LAPD sergeant qualified in tactical waterborne protocols; supervised the 12-person TWB cadre covering 52 miles of Los Angeles waterfront including the Port of LA/Long Beach — third-largest container facility in the world.')

add_role(doc, 'SWAT Officer and Element Leader (Police Officer 3 / 3+1)',
    'Metropolitan Division, Special Weapons and Tactics  •  1996 – 2007')
add_bullet(doc, 'Element member (1996–2003) and element leader (2003–2007). Specializations: sniper, lead climber, tactical waterborne diver, firearms instructor, military and foreign liaison.')

add_role(doc, 'Security Aide to the Chief of Police',
    'Office of the Chief of Police  •  1999 – 2002')
add_bullet(doc, 'Personal protection for Chiefs Bernard C. Parks and Martin Pomeroy. Threat assessments, site surveys, route planning, and liaison with domestic and international law enforcement and government entities.')

add_role(doc, 'Metropolitan Division, C Platoon  /  Field Officer',
    'Southeast and Rampart Patrol Divisions  •  1988 – 1996')
add_bullet(doc, 'Riot response, dignitary protection (U.S. Secret Service / Dept. of State), narcotics, Predator Apprehension Team, Special Problems Unit, and patrol.')

# ── MILITARY SERVICE ─────────────────────────────────────────────
add_section_header(doc, 'Military Service')

add_role(doc, 'Petty Officer, Marine Safety Office',
    'United States Coast Guard Reserve  •  Port of Los Angeles/Long Beach  •  1984 – 1992')
add_bullet(doc, 'Drug interdiction, search and rescue, vessel inspections, hazardous materials investigations, and small arms instruction.')
add_bullet(doc, 'Port Security “A” School, Yorktown, VA (1985); Recruit Training, Cape May, NJ (1984).')
add_bullet(doc, 'National Defense Service Medal; Coast Guard Reserve Good Conduct Medal with Bronze Star Device.')

# ── PROFESSIONAL AFFILIATIONS ────────────────────────────────────
add_section_header(doc, 'Professional Affiliations and Memberships')

affiliations = [
    'International Exchange Alumni, U.S. Department of State (2012 – Present)',
    'Chartered Management Institute, England (2012 – Present)',
    'Royal Society for the Encouragement of Arts, Manufactures and Commerce, England (2011 – Present)',
    'Pi Alpha Alpha – National Honor Society for Public Affairs and Administration (2002 – Present)',
    'University of Southern California Alumni Association (2002 – Present)',
    'American Mensa (2000 – Present)',
    'Professional Association of Diving Instructors, Divemaster #178558 (2000 – Present)',
    'National Tactical Officers Association (1996 – Present)',
    'California Association of Tactical Officers (1996 – Present)',
    'Fraternal Order of Police, Lodge 1 (1990 – Present)',
    'American Legion, Post 291 (2020 – Present)',
]
for a in affiliations:
    add_bullet(doc, a)

# ── HONORS & COMMENDATIONS ───────────────────────────────────────
add_section_header(doc, 'Honors and Commendations')
add_bullet(doc, 'Fulbright Scholarship / Research Fellowship, U.S. Department of State (2011–2012)')
add_bullet(doc, 'Certificate of Merit, Outstanding Master’s Graduate Student, University of Southern California (2002)')
add_bullet(doc, 'Los Angeles Police Department Major Commendations (8)')
add_bullet(doc, 'Los Angeles Police Department Minor Commendations (130+)')

add_running_header(doc, 'Chester Lee McMillion')

docx_filename = 'resume_Chester_McMillion.docx'
doc.save(docx_filename)
print(f'Saved: {docx_filename}')

pdf_path = export_pdf(docx_filename)
if pdf_path:
    print(f'Saved: {pdf_path}')
