"""
URVI Framework — Final Assembly Script
Generates restoration_project/FINAL_DOCUMENT.docx from 07_polished.md per
the Layout Notes and the Corporate / Government style spec in the Work Order.
"""
import os
import re
import sys
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

SOURCE = "restoration_project/07_polished.md"
OUTPUT = "restoration_project/FINAL_DOCUMENT.docx"

DARK_RED = RGBColor(0x8B, 0x00, 0x00)
LIGHT_GRAY = "F2F2F2"
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

class Block:
    def __init__(self, kind, text=None, level=None, items=None, lines=None, meta=None):
        self.kind = kind          # heading | para | bullet | numbered | hr | blockquote | code | layout
        self.text = text
        self.level = level
        self.items = items or []
        self.lines = lines or []
        self.meta = meta or {}


def strip_html_comments(text):
    """Remove <!-- ... --> blocks but keep layout/visual/polish comments out of body text."""
    return re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)


def parse_markdown(md):
    """Parse the polished markdown into a flat list of Block objects."""
    # Remove HTML comments entirely (they're metadata, not body content)
    md = strip_html_comments(md)
    lines = md.split("\n")
    blocks = []
    i = 0

    def flush_para(buf):
        if buf:
            text = " ".join(line.strip() for line in buf).strip()
            if text:
                blocks.append(Block("para", text=text))

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty
        if not stripped:
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            blocks.append(Block("heading", text=m.group(2).strip(), level=level))
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            blocks.append(Block("hr"))
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            j = i + 1
            code_lines = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code_lines.append(lines[j])
                j += 1
            blocks.append(Block("code", lines=code_lines))
            i = j + 1
            continue

        # Blockquote (callout box)
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            blocks.append(Block("blockquote", lines=quote_lines))
            continue

        # Table
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i+1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append(Block("table", lines=table_lines))
            continue

        # Bullet / numbered lists (collect contiguous run)
        if re.match(r"^\s*[-*]\s+", line) or re.match(r"^\s*\d+\.\s+", line):
            list_lines = []
            kind = "bullet" if re.match(r"^\s*[-*]\s+", line) else "numbered"
            while i < len(lines):
                ln = lines[i]
                if re.match(r"^\s*[-*]\s+", ln) or re.match(r"^\s*\d+\.\s+", ln):
                    list_lines.append(ln)
                    i += 1
                elif ln.strip() == "":
                    # blank line ends list
                    break
                elif re.match(r"^\s{2,}\S", ln):
                    # continuation indent — append to last
                    list_lines[-1] += " " + ln.strip()
                    i += 1
                else:
                    break
            items = []
            for ln in list_lines:
                m_b = re.match(r"^\s*[-*]\s+(.*)$", ln)
                m_n = re.match(r"^\s*\d+\.\s+(.*)$", ln)
                if m_b:
                    items.append(("bullet", m_b.group(1).strip()))
                elif m_n:
                    items.append(("numbered", m_n.group(1).strip()))
            blocks.append(Block("list", items=items, meta={"kind": kind}))
            continue

        # Paragraph (gather contiguous non-empty, non-special lines)
        buf = []
        while i < len(lines):
            ln = lines[i]
            s = ln.strip()
            if not s:
                break
            if re.match(r"^(#{1,6})\s+", s) or s == "---" or s.startswith("```") or s.startswith(">") or re.match(r"^\s*[-*]\s+", ln) or re.match(r"^\s*\d+\.\s+", ln):
                break
            if "|" in s and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i+1]):
                break
            buf.append(ln)
            i += 1
        flush_para(buf)

    return blocks


# ---------------------------------------------------------------------------
# Word document construction
# ---------------------------------------------------------------------------

def set_run_font(run, name="Times New Roman", size_pt=11, bold=False, italic=False, color=None, all_caps=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if all_caps:
        run.font.all_caps = True
    # East-Asian font override (Word needs explicit rFonts for non-default fonts)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def add_inline_runs(paragraph, text, base_font="Times New Roman", base_size=11, bold=False, italic=False, color=None):
    """Handle inline **bold**, *italic*, `code` within a paragraph text."""
    # Bold is checked before italic so `**text**` is not parsed as a pair of italics.
    # Italic uses a non-greedy match restricted to a single line so an unmatched
    # asterisk in a long paragraph doesn't swallow content.
    tokens = []
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*\n]+?\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], False, False, False))
        tok = m.group(0)
        if tok.startswith("**"):
            tokens.append((tok[2:-2], True, False, False))
        elif tok.startswith("`"):
            tokens.append((tok[1:-1], False, False, True))
        else:  # single-asterisk italic
            tokens.append((tok[1:-1], False, True, False))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], False, False, False))
    if not tokens:
        tokens = [(text, False, False, False)]

    for tk_text, tk_bold, tk_italic, tk_code in tokens:
        run = paragraph.add_run(tk_text)
        if tk_code:
            set_run_font(run, name="Courier New", size_pt=10, bold=bold or tk_bold, italic=italic or tk_italic, color=color)
        else:
            set_run_font(run, name=base_font, size_pt=base_size, bold=bold or tk_bold, italic=italic or tk_italic, color=color)


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, hex_color="8B0000", style="single", size=8, sides=("top", "left", "bottom", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        if style == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), style)
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), hex_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def clear_cell(cell):
    """Empty a cell's first paragraph and return it for re-use."""
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r.text = ""
    # Remove other paragraphs
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)
    return p


def add_paragraph_border(paragraph, position, size=8, color="8B0000"):
    """Add a single border to a paragraph. position: 'bottom' | 'top' | 'left' | 'right'."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{position}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pBdr.append(el)


def set_page_break_before(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pbb = OxmlElement("w:pageBreakBefore")
    pPr.append(pbb)


def keep_with_next(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def add_toc_field(doc, switches='\\o "1-3" \\h \\z \\u'):
    """Insert a Word auto-TOC field. Word populates it when the doc opens
    (provided updateFields is set to true in settings)."""
    p = doc.add_paragraph()
    run1 = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    fldChar1.set(qn("w:dirty"), "true")
    run1._element.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" TOC {switches} "
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._element.append(fldChar2)

    run4 = p.add_run("Right-click and choose Update Field to populate the Table of Contents.")
    set_run_font(run4, name="Times New Roman", size_pt=11, italic=True)

    run5 = p.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._element.append(fldChar3)


def enable_auto_update_fields(doc):
    """Set <w:updateFields w:val="true"/> in settings.xml so Word refreshes
    fields (TOC, page numbers) when the document is opened."""
    settings = doc.settings.element
    # Remove any existing updateFields to avoid duplicates
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def make_heading(doc, text, level, is_first_h2):
    # Use Word's built-in Heading styles so auto-TOC, navigation, and outline view all work.
    style_name = f"Heading {min(level, 9)}"
    p = doc.add_paragraph(style=style_name)
    pf = p.paragraph_format
    # Spacing: enough to set headings apart visually but not so much it leaves big gaps
    pf.space_before = Pt(14) if level <= 2 else Pt(10)
    pf.space_after = Pt(4)
    # keep_with_next stops a heading from being stranded at the bottom of a page
    keep_with_next(p)

    if level == 1:
        run = p.add_run(text)
        set_run_font(run, name="Times New Roman", size_pt=18, bold=True, color=DARK_RED)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        run = p.add_run(text.upper())
        set_run_font(run, name="Times New Roman", size_pt=14, bold=True, color=BLACK, all_caps=True)
        add_paragraph_border(p, "bottom", size=6, color="8B0000")
        # No forced page break — let content flow naturally
    elif level == 3:
        run = p.add_run(text)
        set_run_font(run, name="Times New Roman", size_pt=12, bold=True, color=BLACK)
    elif level == 4:
        run = p.add_run(text)
        set_run_font(run, name="Times New Roman", size_pt=11, bold=True, italic=True, color=BLACK)
    else:
        run = p.add_run(text)
        set_run_font(run, name="Times New Roman", size_pt=11, bold=True)
    return p


def make_callout_box(doc, lines):
    """Render a blockquote as a callout box: 2-row, 1-column Word table.
    Row 1 (header): dark red fill, white bold caps. Row 2 (body): light gray
    fill, paragraphs preserved. This renders consistently across Word and
    Pages — unlike shaded paragraphs with manual borders, which leave gaps
    on the right edge and break under long text wraps."""
    if not lines:
        return
    header_text = lines[0]
    body_lines = lines[1:]
    header_clean = re.sub(r"^\*\*(.+?)\*\*$", r"\1", header_text.strip())

    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Width: full content area (6.5" given Letter + 1" margins)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5)

    # ---- Header row ----
    hcell = table.rows[0].cells[0]
    set_cell_shading(hcell, "8B0000")
    set_cell_borders(hcell, "8B0000", "single", 12)
    hcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hp = clear_cell(hcell)
    hp.paragraph_format.space_before = Pt(3)
    hp.paragraph_format.space_after = Pt(3)
    hp.paragraph_format.line_spacing = 1.0
    hrun = hp.add_run(header_clean.upper())
    set_run_font(hrun, name="Times New Roman", size_pt=11, bold=True, color=WHITE, all_caps=True)

    # ---- Body row ----
    bcell = table.rows[1].cells[0]
    set_cell_shading(bcell, LIGHT_GRAY)
    set_cell_borders(bcell, "8B0000", "single", 12)
    # First body paragraph: reuse the cell's existing one
    first_par = clear_cell(bcell)
    for idx, raw in enumerate(body_lines):
        if idx == 0:
            bp = first_par
        else:
            bp = bcell.add_paragraph()
        bp.paragraph_format.space_before = Pt(3) if idx == 0 else Pt(2)
        bp.paragraph_format.space_after = Pt(3) if idx == len(body_lines) - 1 else Pt(2)
        bp.paragraph_format.line_spacing = 1.15
        add_inline_runs(bp, raw, base_size=11)

    # Spacer after the table so the next paragraph isn't flush against the border
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


TIGHT_LIST_THRESHOLD = 6  # Lists at or below this many items are kept together on one page


def make_list(doc, items):
    """Render a list of (kind, text) tuples. Tight lists (≤ TIGHT_LIST_THRESHOLD
    items) get keep_with_next on every item except the last, so Word treats the
    whole list as a single block and won't break it across pages. Longer lists
    are allowed to flow naturally."""
    tight = len(items) <= TIGHT_LIST_THRESHOLD
    last_idx = len(items) - 1
    for i, (kind, txt) in enumerate(items):
        # Checklist items: text starts with ☐
        if txt.startswith("☐"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run("☐ ")
            set_run_font(run, name="Times New Roman", size_pt=11)
            remainder = txt[1:].lstrip()
            add_inline_runs(p, remainder, base_size=11)
            if tight and i < last_idx:
                keep_with_next(p)
            continue

        p = doc.add_paragraph(style="List Bullet" if kind == "bullet" else "List Number")
        p.paragraph_format.space_after = Pt(3)
        for r in list(p.runs):
            r.text = ""
        add_inline_runs(p, txt, base_size=11)
        if tight and i < last_idx:
            keep_with_next(p)


def make_paragraph(doc, text):
    # Special handling: Phase transition labels like "**PHASE 1 → PHASE 2:**"
    if text.startswith("**PHASE") and text.endswith("**"):
        clean = text.strip("*")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)
        keep_with_next(p)
        set_run_font(p.add_run(clean), name="Times New Roman", size_pt=11, bold=True, color=DARK_RED)
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    add_inline_runs(p, text, base_size=11)


def make_code_block(doc, lines):
    """Render a code block. The URVI Figure 1 ASCII org chart gets special
    treatment — detected by the presence of 'Unified Command' + 'Contact Team' —
    and rendered as a structured table org chart instead."""
    joined = "\n".join(lines)
    if "Unified Command" in joined and "Contact Team" in joined and "RTF" in joined:
        make_urvi_org_chart(doc)
        return
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(ln), name="Courier New", size_pt=9)


def _box_cell(cell, title, subtitle=None, font_size=10, subtitle_size=9):
    """Style a cell as an org-chart box: light gray fill, dark red border,
    centered bold title with optional subtitle line."""
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_borders(cell, "8B0000", "single", 8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2 if subtitle else 4)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run(title), name="Times New Roman", size_pt=font_size, bold=True)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.line_spacing = 1.0
        set_run_font(p2.add_run(subtitle), name="Times New Roman", size_pt=subtitle_size, italic=False)


def _connector_cell(cell):
    """Style a cell as a no-border, no-fill cell with a centered vertical bar."""
    set_cell_borders(cell, "FFFFFF", "nil")
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run("│"), name="Times New Roman", size_pt=12, bold=True, color=DARK_RED)


def _blank_cell(cell):
    set_cell_borders(cell, "FFFFFF", "nil")
    clear_cell(cell)


def make_urvi_org_chart(doc):
    """Render the URVI command-structure org chart as a 5-row × 7-col table.
    Spacer columns (1, 3, 5) give each node its own visible box with
    whitespace between siblings — adjacent shaded cells with shared borders
    would otherwise merge visually into one wide compartmentalized block.

    Grid layout (B = box, S = spacer, V = vertical connector):
        Col:  0   1   2   3   4   5   6
        R0:   ┌────── Unified Command ──────┐
        R1:   ┌─────────── │ ───────────────┐
        R2:   ┌── Contact ──┐ S ┌── Rescue ──┐
        R3:   V   .   V   .   V   .   V
        R4:   B   .   B   .   B   .   B
    """
    table = doc.add_table(rows=5, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths — box columns are wider than spacer columns
    # 4 boxes × 1.30" + 3 spacers × 0.43" = 5.20 + 1.29 = 6.49" (full content width)
    BOX_W = 1.30
    SPACER_W = 0.43
    col_widths_in = [BOX_W, SPACER_W, BOX_W, SPACER_W, BOX_W, SPACER_W, BOX_W]

    # Set grid widths so Word and Pages respect the layout
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    gridCols = tblGrid.findall(qn("w:gridCol"))
    for col, w_in in zip(gridCols, col_widths_in):
        col.set(qn("w:w"), str(int(w_in * 1440)))  # inches → twips (1 inch = 1440)
        col.set(qn("w:type"), "dxa")

    # Also set each cell's preferred width
    for row in table.rows:
        for cell, w_in in zip(row.cells, col_widths_in):
            cell.width = Inches(w_in)

    # Row 0 — Unified Command (merged across all 7 columns)
    r0 = table.rows[0]
    top = r0.cells[0].merge(r0.cells[6])
    _box_cell(top, "Unified Command", subtitle="(Law Enforcement / Fire / EMS)")

    # Row 1 — vertical connector below Unified Command (merged across all 7)
    r1 = table.rows[1]
    conn1 = r1.cells[0].merge(r1.cells[6])
    _connector_cell(conn1)

    # Row 2 — Contact Group (cols 0-2), spacer (col 3), Rescue Group (cols 4-6)
    r2 = table.rows[2]
    contact = r2.cells[0].merge(r2.cells[2])
    _box_cell(contact, "Contact Group", subtitle="(Supervisor)")
    _blank_cell(r2.cells[3])
    rescue = r2.cells[4].merge(r2.cells[6])
    _box_cell(rescue, "Rescue Group", subtitle="(Supervisor)")

    # Row 3 — vertical connectors above each leaf box; spacers in between
    r3 = table.rows[3]
    for idx, c in enumerate(r3.cells):
        if idx in (0, 2, 4, 6):
            _connector_cell(c)
        else:
            _blank_cell(c)

    # Row 4 — four leaf boxes (CT1, CT2, RTF 1, RTF 2) separated by spacers
    r4 = table.rows[4]
    labels = ["Contact Team 1", "Contact Team 2", "RTF 1", "RTF 2"]
    label_iter = iter(labels)
    for idx, c in enumerate(r4.cells):
        if idx in (0, 2, 4, 6):
            _box_cell(c, next(label_iter), subtitle=None)
        else:
            _blank_cell(c)

    # Trailing spacer paragraph so the caption isn't flush against the chart
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def make_table(doc, lines):
    """Parse and render a markdown table."""
    rows = []
    for ln in lines:
        if re.match(r"^\s*\|?\s*:?-+", ln):  # separator row
            continue
        # Trim leading/trailing pipes
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        # Pad short rows
        cells = row + [""] * (ncols - len(row))
        for j, cell in enumerate(cells):
            tc = table.rows[i].cells[j]
            tc.text = ""
            p = tc.paragraphs[0]
            if i == 0:
                # Header row: bold
                add_inline_runs(p, cell, base_size=10, bold=True)
            else:
                add_inline_runs(p, cell, base_size=10)


# ---------------------------------------------------------------------------
# Header + footer
# ---------------------------------------------------------------------------

def add_page_number(paragraph):
    """Insert a PAGE field into the given paragraph (right side after a tab)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    set_run_font(run, name="Times New Roman", size_pt=10)


def set_running_header(section, doc_title):
    section.different_first_page_header_footer = True
    # First-page header stays blank
    first_header = section.first_page_header
    for p in first_header.paragraphs:
        for r in list(p.runs):
            r.text = ""
    # Regular header on pages 2+
    header = section.header
    header.is_linked_to_previous = False
    # Use first paragraph
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Clear any existing content
    for r in list(p.runs):
        r.text = ""
    # Add title on left, tab, page number on right
    pPr = p._element.get_or_add_pPr()
    # Add tab stops: center and right (right at 6.5" given 1" margins on Letter)
    tabs = OxmlElement("w:tabs")
    tab_right = OxmlElement("w:tab")
    tab_right.set(qn("w:val"), "right")
    tab_right.set(qn("w:pos"), "9360")  # twips = 6.5 inches
    tabs.append(tab_right)
    pPr.append(tabs)
    run = p.add_run(doc_title)
    set_run_font(run, name="Times New Roman", size_pt=10)
    p.add_run("\t")
    add_page_number(p)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def build_cover_page(doc):
    # The cover page is the first section. Add several blank paragraphs to push title down a bit.
    for _ in range(6):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    set_run_font(p_title.add_run("UNIFIED RESPONSE TO VIOLENT INCIDENTS (URVI)"), name="Times New Roman", size_pt=18, bold=True, color=DARK_RED)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    set_run_font(p_sub.add_run("Framework for Los Angeles County Operational Guidelines"), name="Times New Roman", size_pt=13, italic=True, color=BLACK)

    p_prep = doc.add_paragraph()
    p_prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p_prep.add_run("Prepared by the URVI Guidelines Committee"), name="Times New Roman", size_pt=11)

    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org.paragraph_format.space_after = Pt(0)
    set_run_font(p_org.add_run("Los Angeles County Public Safety Agencies"), name="Times New Roman", size_pt=11)

    # Vertical gap
    for _ in range(8):
        doc.add_paragraph()

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p_ver.add_run("Version 1.0 | Effective: [DATE]"), name="Times New Roman", size_pt=10)


# ---------------------------------------------------------------------------
# Main assembly
# ---------------------------------------------------------------------------

def main():
    with open(SOURCE, "r", encoding="utf-8") as f:
        md = f.read()

    # Stop before the "## Completion Checklist" — that's audit/polish metadata, not body content.
    # Also stop before "## Layout Notes for Final Assembly" — those are instructions for THIS step.
    cutoff_idx = md.find("\n## Layout Notes for Final Assembly")
    if cutoff_idx == -1:
        cutoff_idx = md.find("\n## Completion Checklist")
    if cutoff_idx != -1:
        md = md[:cutoff_idx]

    blocks = parse_markdown(md)

    doc = Document()

    # Configure section: Letter, 1" margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    # ----- Cover Page -----
    build_cover_page(doc)
    add_page_break(doc)

    # ----- Body -----
    # Running header on pages 2+, none on cover (different_first_page_header_footer)
    set_running_header(section, "Unified Response to Violent Incidents (URVI)")

    h2_count = 0
    skip_first_h1 = False
    in_plaintext_toc = False  # When true, suppress blocks (we replaced TOC with a field)
    body_started = False  # Suppress source's cover-duplicate paragraph before first H2
    last_heading_text = None  # Used to drop a paragraph that just re-states the heading

    def next_visible_block(start_idx):
        """Return the next block from start_idx onward that would actually
        render. Used to peek for tight-list followers."""
        for j in range(start_idx, len(blocks)):
            b = blocks[j]
            if b.kind == "hr":
                continue
            return b
        return None

    for idx, blk in enumerate(blocks):
        if blk.kind == "heading":
            last_heading_text = blk.text.strip().upper()
            if blk.level == 1:
                # The first H1 is the doc title — already rendered as cover; skip it.
                if not skip_first_h1:
                    skip_first_h1 = True
                    continue
                in_plaintext_toc = False
                make_heading(doc, blk.text, 1, is_first_h2=False)
            elif blk.level == 2:
                # Detect numbered major sections (e.g., "1. EXECUTIVE SUMMARY", "TABLE OF CONTENTS", "FIGURE 1: ...")
                in_plaintext_toc = False
                body_started = True
                is_first_h2 = (h2_count == 0)
                make_heading(doc, blk.text, 2, is_first_h2=is_first_h2)
                h2_count += 1
                if blk.text.strip().upper() == "TABLE OF CONTENTS":
                    # Insert a Word auto-TOC field and force a page break afterward so
                    # the first body section starts at the top of a fresh page. This
                    # is the only forced page break inside the body — section flow
                    # otherwise pages naturally so short sections don't strand whitespace.
                    add_toc_field(doc)
                    add_page_break(doc)
                    in_plaintext_toc = True
            else:
                make_heading(doc, blk.text, blk.level, is_first_h2=False)
        elif not body_started:
            # Everything before the first H2 is cover-duplicate content (subtitle,
            # "Prepared by...", agency line) — already rendered on the cover page.
            continue
        elif in_plaintext_toc:
            # Skip any blocks that are part of the original plain-text TOC
            continue
        elif blk.kind == "para":
            # Drop a paragraph that just re-states the heading immediately above
            # (e.g., "**FIGURE 1: EXAMPLE BASIC URVI ORGANIZATION**" right after
            # the H2 of the same name in the source).
            stripped = blk.text.strip().strip("*").upper()
            if last_heading_text and stripped == last_heading_text:
                continue
            make_paragraph(doc, blk.text)
            # If this paragraph introduces a tight list (ends with ":" and the
            # next block is a list of ≤ TIGHT_LIST_THRESHOLD items), bind it to
            # that list so Word keeps the lead-in on the same page as its items.
            if blk.text.rstrip().endswith(":"):
                nxt = next_visible_block(idx + 1)
                if nxt is not None and nxt.kind == "list" and len(nxt.items) <= TIGHT_LIST_THRESHOLD:
                    keep_with_next(doc.paragraphs[-1])
        elif blk.kind == "list":
            make_list(doc, blk.items)
        elif blk.kind == "blockquote":
            make_callout_box(doc, blk.lines)
        elif blk.kind == "code":
            make_code_block(doc, blk.lines)
        elif blk.kind == "table":
            make_table(doc, blk.lines)
        elif blk.kind == "hr":
            # Treat as visual separator — small spacing paragraph; suppress consecutive HRs
            continue

    # ----- Auto-update fields on open (populates the TOC field) -----
    enable_auto_update_fields(doc)

    # ----- Save -----
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

    # ----- Self-verify -----
    assert os.path.exists(OUTPUT), "File not created"
    size = os.path.getsize(OUTPUT)
    assert size > 5000, f"File too small: {size} bytes"
    doc2 = Document(OUTPUT)
    headings = [p for p in doc2.paragraphs if p.style.name.startswith("Heading")]
    print(f"File size: {size:,} bytes")
    print(f"Total paragraphs: {len(doc2.paragraphs)}")
    print(f"Headings detected (built-in style): {len(headings)}")
    print(f"Tables: {len(doc2.tables)}")

    # Word counts for sanity
    total_text = "\n".join(p.text for p in doc2.paragraphs)
    word_count = len(total_text.split())
    print(f"Estimated word count: {word_count}")
    # Estimate pages ~ 350 words/page for Times New Roman 11pt 1.15 spacing
    est_pages = max(1, (word_count // 350) + 4)  # +cover, +TOC, +figure, +section page breaks
    print(f"Estimated pages: ~{est_pages}")

    # ----- PDF export -----
    pdf_path = os.path.abspath(OUTPUT.replace(".docx", ".pdf"))
    docx_abs = os.path.abspath(OUTPUT)
    if export_pdf_via_word(docx_abs, pdf_path):
        print(f"PDF generated: {pdf_path}")
    else:
        print("PDF conversion not available — open the .docx in Word and use File → Save As → PDF.")


def export_pdf_via_word(docx_abs, pdf_abs):
    """Convert .docx → .pdf. Tries (1) Word + update-fields, (2) Pages export.
    Falls back gracefully if both fail."""
    import subprocess

    # Step 1: open the doc in Word and update fields so the TOC is populated.
    # This Word version's AppleScript dictionary rejects `save as PDF` but
    # accepts `update fields` and `save`. We refresh fields and re-save the
    # docx, then hand off to a different tool for the PDF.
    refresh_script = f'''
    tell application "Microsoft Word"
        activate
        open file name "{docx_abs}"
        set theDoc to active document
        update fields theDoc
        save theDoc
        close theDoc saving no
    end tell
    '''
    try:
        result = subprocess.run(
            ["osascript", "-e", refresh_script],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            print(f"Word field refresh stderr: {result.stderr.strip()}")
        else:
            print("Word field refresh: TOC populated and .docx re-saved.")
    except Exception as e:
        print(f"Word refresh failed: {e}")

    # Step 2: PDF export via Pages (well-supported AppleScript export to PDF)
    pages_script = f'''
    tell application "Pages"
        activate
        set theDoc to open (POSIX file "{docx_abs}")
        delay 2
        export theDoc to (POSIX file "{pdf_abs}") as PDF
        close theDoc saving no
    end tell
    '''
    try:
        result = subprocess.run(
            ["osascript", "-e", pages_script],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and os.path.exists(pdf_abs):
            print("PDF exported via Pages.")
            return True
        print(f"Pages export stderr: {result.stderr.strip()}")
    except Exception as e:
        print(f"Pages export failed: {e}")

    return False


if __name__ == "__main__":
    main()
